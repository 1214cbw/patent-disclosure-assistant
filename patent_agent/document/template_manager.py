from __future__ import annotations

from pathlib import Path
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm

from .styles import configure_styles, set_run_font


def _page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar"); begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText"); instruction.set(qn("xml:space"), "preserve"); instruction.text = " PAGE "
    end = OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instruction, end))


class TemplateManager:
    def __init__(self, template_root: Path):
        self.root = Path(template_root)

    @property
    def default_disclosure(self) -> Path:
        return self.root / "disclosure" / "default_cn_disclosure.docx"

    def ensure_default(self) -> Path:
        path = self.default_disclosure
        if path.exists(): return path
        path.parent.mkdir(parents=True, exist_ok=True)
        document = configure_styles(Document())
        section = document.sections[0]
        section.top_margin = Cm(2.54); section.bottom_margin = Cm(2.54); section.left_margin = Cm(2.8); section.right_margin = Cm(2.6)
        footer = section.footer.paragraphs[0]
        _page_number(footer)
        document.core_properties.title = "中国专利技术交底书模板"
        document.core_properties.subject = "Patent Agent V1 default disclosure template"
        document.save(path)
        return path

