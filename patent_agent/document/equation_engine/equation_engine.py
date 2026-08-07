"""Public Patent Agent-facing API for native, editable Word mathematics."""
from __future__ import annotations

import logging
from pathlib import Path
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Inches
from .omml_renderer import LatexParseError, latex_to_omml

LOG = logging.getLogger(__name__)


def _set_cell_borders_none(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = OxmlElement(f"w:{edge}"); node.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val", "nil"); borders.append(node)
    tbl_pr.append(borders)


class EquationEngine:
    """Insert OMML; callers never have to handle XML, XSLT, or COM."""
    def to_omml(self, latex: str, *, display: bool = False):
        return latex_to_omml(latex, display=display)

    def insert_inline(self, paragraph, latex: str):
        omml = self.to_omml(latex)
        paragraph._p.append(omml)
        return omml

    def insert_display(self, document, latex: str, *, label: str | None = None):
        if label:
            document.add_paragraph(label)
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        omml = self.to_omml(latex, display=True)
        paragraph._p.append(omml)
        return paragraph

    def insert_numbered(self, document, latex: str, number: int):
        table = document.add_table(rows=1, cols=3)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        _set_cell_borders_none(table)
        widths = (Inches(1.0), Inches(4.8), Inches(1.0))
        for cell, width in zip(table.rows[0].cells, widths): cell.width = width
        equation_paragraph = table.cell(0, 1).paragraphs[0]
        equation_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        equation_paragraph._p.append(self.to_omml(latex, display=True))
        number_paragraph = table.cell(0, 2).paragraphs[0]
        number_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        number_paragraph.add_run(f"({number})")
        return table

    @staticmethod
    def word_com_available() -> bool:
        try:
            import win32com.client  # noqa: F401
            return True
        except ImportError:
            return False

    def validate_with_word(self, docx_path: str | Path, export_pdf: bool = True) -> dict:
        """Open with Word and count OMaths.  This is also the fallback channel.

        Unsupported future syntax can be rendered by Word linear-math input in a
        dedicated fallback builder without changing the public API.
        """
        if not self.word_com_available():
            return {"available": False, "reason": "pywin32 is not installed"}
        import win32com.client
        word = None; document = None
        path = Path(docx_path).resolve()
        try:
            word = win32com.client.DispatchEx("Word.Application")
            word.Visible = False; word.DisplayAlerts = 0
            document = word.Documents.Open(str(path), ReadOnly=False, AddToRecentFiles=False)
            count = int(document.OMaths.Count)
            document.Save()
            pdf = path.with_suffix(".pdf")
            if export_pdf:
                document.ExportAsFixedFormat(str(pdf), 17)
            return {"available": True, "omaths_count": count, "pdf": str(pdf) if pdf.exists() else None}
        except Exception as exc:
            return {"available": True, "error": f"{type(exc).__name__}: {exc}"}
        finally:
            if document is not None: document.Close(SaveChanges=False)
            if word is not None: word.Quit()

