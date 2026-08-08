from __future__ import annotations

import re
from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

from patent_agent.core.patent_ast import PatentDocumentAST, PatentNode
from .equation_engine import EquationEngine
from .styles import configure_styles, set_run_font
from .template_manager import TemplateManager


class DocumentRenderer:
    """Deterministic AST-to-DOCX renderer with inline OMML math support."""
    def __init__(self, template_root: Path):
        self.templates = TemplateManager(template_root)
        self.equations = EquationEngine()
        # Lazy-init math detector
        self._math_detector = None

    @property
    def math_detector(self):
        if self._math_detector is None:
            from .math_detector import MathSpanDetector
            self._math_detector = MathSpanDetector()
        return self._math_detector

    def render(self, ast: PatentDocumentAST, output_path: Path, template_path: Path | None = None) -> Path:
        template = template_path or self.templates.ensure_default()
        document = configure_styles(Document(template))
        document.core_properties.title = ast.title
        document.core_properties.subject = f"Patent Agent {ast.kind}"
        for node in ast.nodes:
            self._render_node(document, node)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        document.save(output_path)
        return output_path

    def _render_node(self, document, node: PatentNode):
        if node.type == "heading":
            paragraph = document.add_paragraph(style="Title" if node.level == 0 else f"Heading {min(node.level or 1, 2)}")
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if node.level == 0 else WD_ALIGN_PARAGRAPH.LEFT
            set_run_font(paragraph.add_run(node.value), east_asia="黑体", size=20 if node.level == 0 else 14, bold=True)
        elif node.type == "paragraph":
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.first_line_indent = Cm(0.74)
            if node.children:
                for child in node.children: self._render_inline(paragraph, child)
            else:
                # Use MathSpanDetector to render text with inline OMML
                self._render_text_with_math(paragraph, node.value)
        elif node.type == "display_equation":
            if node.number is not None: self.equations.insert_numbered(document, node.latex, node.number)
            else: self.equations.insert_display(document, node.latex)
        elif node.type == "figure":
            paragraph = document.add_paragraph(); paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # V6.6: aspect-ratio-aware embed size; never exceed usable page
            width, height = self._figure_embed_size(node.path)
            paragraph.paragraph_format.keep_with_next = True
            paragraph.paragraph_format.keep_together = True
            paragraph.add_run().add_picture(node.path, width=width, height=height)
            caption = document.add_paragraph(style="Patent Caption"); caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption.paragraph_format.keep_with_next = False
            caption.paragraph_format.keep_together = True
            # Strip duplicate "图N" prefix from value if present
            import re
            figure_title = node.value or ""
            figure_title = re.sub(r'^图\s*\d+\s*[：:.\s]*', '', figure_title).strip()
            set_run_font(caption.add_run(f"图{node.number}  {figure_title}"), size=10.5)
        elif node.type == "list":
            for child in node.children:
                paragraph = document.add_paragraph(style="List Bullet"); set_run_font(paragraph.add_run(child.value))
        elif node.type == "claim":
            paragraph = document.add_paragraph(); paragraph.paragraph_format.first_line_indent = Cm(0.74)
            set_run_font(paragraph.add_run(f"{node.number}. {node.value}"))
        elif node.type == "inventor_question":
            paragraph = document.add_paragraph(); set_run_font(paragraph.add_run("[待发明人确认] " + node.value), bold=True)
        elif node.type == "page_break": document.add_page_break()

    def _render_inline(self, paragraph, node: PatentNode):
        if node.type == "text": set_run_font(paragraph.add_run(node.value))
        elif node.type == "inline_math": self.equations.insert_inline(paragraph, node.latex)
        elif node.type == "equation_reference": set_run_font(paragraph.add_run(f"式（{self._ref_number(node.target)}）"))
        elif node.type == "figure_reference": set_run_font(paragraph.add_run(f"图{self._ref_number(node.target)}"))
        elif node.type == "source_citation": set_run_font(paragraph.add_run(node.value), size=9)

    def _render_text_with_math(self, paragraph, text: str):
        """Render paragraph text with automatic inline OMML math detection.

        Scans the text for registered math symbols and renders them
        as proper OMML inline equations, while keeping Chinese text as-is.
        """
        if not text.strip():
            return

        spans = self.math_detector.convert_paragraph(text)

        for span in spans:
            if span["type"] == "text":
                set_run_font(paragraph.add_run(span["value"]))
            elif span["type"] == "inline_math":
                try:
                    self.equations.insert_inline(paragraph, span["latex"])
                except Exception:
                    # Fallback: render as italic text
                    run = paragraph.add_run(span["latex"])
                    run.italic = True
                    set_run_font(run)

    @staticmethod
    def _ref_number(target: str) -> int:
        digits = "".join(ch for ch in target if ch.isdigit())
        return int(digits or "0")

    @staticmethod
    def _figure_embed_size(image_path: str):
        """Compute embed size (cm) preserving aspect ratio.

        Max width 13.5cm; max height limited to usable A4 page height so a
        tall figure is never silently clipped by Word (the 'looks cropped'
        failure mode). Returns (width_cm, height_cm).
        """
        from PIL import Image as PILImage
        MAX_W_CM = 13.5
        MAX_H_CM = 23.0  # A4 29.7cm - 2*~2.5cm margins - caption space
        try:
            with PILImage.open(image_path) as im:
                iw, ih = im.size
        except Exception:
            return Cm(13.5), None
        if iw <= 0 or ih <= 0:
            return Cm(13.5), None
        aspect = ih / iw
        w_cm = MAX_W_CM
        h_cm = w_cm * aspect
        if h_cm > MAX_H_CM:
            h_cm = MAX_H_CM
            w_cm = h_cm / aspect
        return Cm(max(4.0, w_cm)), Cm(max(3.0, h_cm))

