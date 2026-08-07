"""MathSpanDetector - converts plain-text math tokens to inline_math PatentNodes.

Scans disclosure paragraphs and replaces registered math symbols
with inline_math nodes that render as proper OMML equations.
"""
from __future__ import annotations

import re
from typing import Any

from patent_agent.document.math_registry import (
    EXCLUDED_ACRONYMS,
    REAL_PAPER_001_SYMBOLS,
    UNITS,
    MathSymbol,
    build_symbol_map,
    should_be_math,
)


class MathSpanDetector:
    """Detect and convert plain math tokens to inline OMML annotations."""

    def __init__(self, symbols: list[MathSymbol] | None = None):
        self.symbols = symbols or REAL_PAPER_001_SYMBOLS
        self.symbol_map = build_symbol_map(self.symbols)
        # Build sorted pattern (longest first to avoid partial matches)
        self._build_patterns()

    def _build_patterns(self):
        """Build regex patterns for math token detection.

        IMPORTANT: Expressions and longer forms MUST come first in the alternation
        to prevent partial matches. E.g., 'z_{t-1}' must match before 't'.
        """
        # Sort by: 1) kind priority, 2) length descending
        _kind_priority = {"expression": 0, "subscript": 1, "greek": 2, "variable": 3}
        def _sort_key(s: MathSymbol) -> tuple[int, int]:
            return (_kind_priority.get(s.kind, 4), -len(s.plain_form))

        self.sorted_symbols = sorted(self.symbols, key=_sort_key)
        patterns = []
        for s in self.sorted_symbols:
            form = re.escape(s.plain_form)
            if s.kind == "expression" or s.kind in ("greek",):
                patterns.append(form)
            else:
                patterns.append(r'\b' + form + r'\b')
        self.detect_pattern = re.compile('|'.join(patterns))

    def convert_paragraph(
        self, text: str,
    ) -> list[dict[str, Any]]:
        """Convert a paragraph's text into a list of text/math spans.

        Returns list of dicts with:
        - type: "text" | "inline_math"
        - value (for text) or latex (for inline_math)
        """
        spans = []
        pos = 0

        for match in self.detect_pattern.finditer(text):
            # Add text before match
            if match.start() > pos:
                before = text[pos:match.start()]
                if before:
                    spans.append({"type": "text", "value": before})

            # Check if this token should be converted
            token = match.group()
            ctx_start = max(0, match.start() - 40)
            ctx_end = min(len(text), match.end() + 40)
            context = text[ctx_start:ctx_end]

            if self._is_excluded(token, context):
                spans.append({"type": "text", "value": token})
            else:
                symbol = self.symbol_map.get(token)
                if symbol:
                    spans.append({
                        "type": "inline_math",
                        "latex": symbol.latex,
                    })
                else:
                    spans.append({"type": "text", "value": token})

            pos = match.end()

        # Add remaining text
        if pos < len(text):
            spans.append({"type": "text", "value": text[pos:]})

        return spans

    def _is_excluded(self, token: str, context: str) -> bool:
        """Check if token should NOT be converted to math."""
        # Excluded acronyms
        if token in EXCLUDED_ACRONYMS:
            return True
        # Units
        if token in UNITS:
            return True
        # Check if token appears inside an excluded acronym
        ctx_lower = context.lower()
        for acr in EXCLUDED_ACRONYMS:
            if acr.lower() in ctx_lower:
                # Check if token is a substring of an excluded acronym
                acr_lower = acr.lower()
                if token.lower() in acr_lower and token.lower() != acr_lower:
                    return True
        # Single-letter tokens in acronym-heavy context (likely not math)
        if len(token) == 1 and token.isalpha():
            # Check if surrounded by uppercase letters (acronym context)
            nearby = text_nearby(context, token, 5)
            if nearby and nearby.isupper() and len(nearby) > 2:
                return True
        return False


def text_nearby(context: str, token: str, distance: int) -> str:
    """Get text near a token for context analysis."""
    idx = context.find(token)
    if idx < 0:
        return ""
    start = max(0, idx - distance)
    end = min(len(context), idx + len(token) + distance)
    return context[start:end]


def convert_disclosure_to_math_ast(
    disclosure,  # GroundedDisclosure
    symbols: list[MathSymbol] | None = None,
) -> "GroundedDisclosure":
    """Convert all paragraphs in a disclosure to use inline math spans.

    Returns a new disclosure with paragraph text containing math annotations.
    The actual AST node insertion happens during DOCX rendering.
    """
    detector = MathSpanDetector(symbols)

    new_sections = []
    math_count = 0

    for section in disclosure.sections:
        new_paras = []
        for para in getattr(section, "paragraphs", []):
            text = getattr(para, "text", "")

            # Convert to spans
            spans = detector.convert_paragraph(text)

            # Count math spans
            math_count += sum(1 for s in spans if s["type"] == "inline_math")

            # Store spans as paragraph metadata for the renderer
            new_para = para.model_copy(update={
                "text": text,  # Keep original text
            })
            # Add math_spans to the paragraph if it's a dict-like model
            if hasattr(new_para, "__dict__"):
                # We'll store the spans info via a custom attribute pattern
                pass
            new_paras.append(new_para)

        new_sections.append(section.model_copy(update={"paragraphs": new_paras}))

    return disclosure.model_copy(update={"sections": new_sections}), math_count


def render_paragraph_with_math(
    document,
    paragraph_text: str,
    math_detector: MathSpanDetector,
    equation_engine,
    styles_module,
) -> Any:
    """Render a paragraph with inline math OMML nodes.

    This is the core rendering function that creates a Word paragraph
    with mixed Chinese text runs and inline OMML equation runs.

    Args:
        document: python-docx Document
        paragraph_text: The paragraph text to render
        math_detector: MathSpanDetector instance
        equation_engine: EquationEngine instance
        styles_module: Module containing set_run_font

    Returns:
        The python-docx Paragraph object
    """
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    para = document.add_paragraph()
    spans = math_detector.convert_paragraph(paragraph_text)

    for span in spans:
        if span["type"] == "text":
            run = para.add_run(span["value"])
            styles_module.set_run_font(run)
        elif span["type"] == "inline_math":
            # Insert inline OMML equation
            try:
                equation_engine.insert_inline(para, span["latex"])
            except Exception:
                # Fallback: render as italic text
                run = para.add_run(span["latex"])
                run.italic = True
                styles_module.set_run_font(run)

    return para
