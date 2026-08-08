"""V6.7 disclosure figure final polish - 9 named tests.

Covers: fig-2 crop removes surrounding prose / original English caption /
preserves annotations; generated figures carry no internal titles;
BlankPageValidator flags empty pages; no forced page break after captions;
fig-3 parameter bridge is dashed (visually distinct); fig-4 latent/decoder/
output semantic split; Word captions consistent (centered, uniform format).
"""
from __future__ import annotations

import json
from pathlib import Path

import pytest

from patent_agent.agents.figure_planner import FigurePlanner
from patent_agent.core.models import DisclosureDraft, FigureSpec, TechnicalUnderstandingResult
from patent_agent.core.patent_ast import PatentDocumentAST, PatentNode
from patent_agent.document import DocumentRenderer
from patent_agent.document.ast_factory import disclosure_to_ast
from patent_agent.document.blank_page_validator import BlankPageValidator
from patent_agent.document.figure_renderer import PatentFigureRenderer, _layout_report_path
from patent_agent.document.source_figure_cropper import SourceFigureContentCropper

SRC_PDF = Path("workspace/private_cases/REAL-PAPER-001/source/A Motor Topology Image Generation Method Based on Latent Diffusion Model.pdf")
FIG2_V67 = Path("workspace/private_cases/REAL-PAPER-001/extracted_figures/fig2_design_variables_v67.png")
GOLDEN_KEY = "REAL-PAPER-001:2:fig2"


def _plan_motor_figures() -> list[FigureSpec]:
    u = TechnicalUnderstandingResult.model_construct(
        steps=["motor rotor topology design", "latent diffusion training", "generation"],
        components=["motor", "rotor", "latent diffusion", "topology"],
    )
    return FigurePlanner().from_understanding(u, case_id="REAL-PAPER-001")


def _render(figure: FigureSpec, tmp_path: Path) -> dict:
    PatentFigureRenderer().render(figure, tmp_path)
    return json.loads(_layout_report_path(tmp_path, figure.number).read_text(encoding="utf-8"))


# ── 1. source_figure_remove_surrounding_text_test ────────────────

def test_source_figure_remove_surrounding_text_test(tmp_path: Path):
    """Fig-2 crop must NOT contain body-prose residue above the figure."""
    cropper = SourceFigureContentCropper(SRC_PDF)
    analysis = cropper.analyze_page(1, [55.0, 380.0, 300.0, 540.0])
    assert analysis["prose_ink_bottom_pt"] is not None
    bbox, method = cropper.figure_bbox(1, [55.0, 380.0, 300.0, 540.0], 6.0)
    assert method == "content_bbox"
    assert bbox[1] > analysis["prose_ink_bottom_pt"], \
        "crop 顶部必须位于正文文字之下（移除正文残句）"
    assert bbox[3] > bbox[1], "裁剪矩形无效（inverted）"
    # pixel proof: the crop top band carries no text-like ink
    from PIL import Image
    assert FIG2_V67.exists(), "v67 crop 文件必须存在"
    with Image.open(FIG2_V67).convert("L") as im:
        w, h = im.size
        dark = sum(1 for x in range(0, w, 4) for y in range(0, h // 8, 4)
                   if im.getpixel((x, y)) < 128)
        assert dark / (len(range(0, w, 4)) * len(range(0, h // 8, 4))) < 0.05, \
            "crop 顶部带存在疑似正文文字墨迹"


# ── 2. source_figure_remove_original_caption_test ────────────────

def test_source_figure_remove_original_caption_test():
    """Fig-2 crop must exclude the original English caption."""
    cropper = SourceFigureContentCropper(SRC_PDF)
    analysis = cropper.analyze_page(1, [55.0, 380.0, 300.0, 540.0])
    assert analysis["caption_bbox_pt"], "原图英文 caption 应被识别"
    cap_y0 = min(c[1] for c in analysis["caption_bbox_pt"])
    bbox, _ = cropper.figure_bbox(1, [55.0, 380.0, 300.0, 540.0], 6.0)
    assert bbox[3] < cap_y0, "crop 底部侵入原始英文 caption"


# ── 3. source_figure_preserve_annotations_test ───────────────────

def test_source_figure_preserve_annotations_test():
    """Fig-2 crop must keep the in-figure label (hbs1) and geometry."""
    from PIL import Image
    assert FIG2_V67.exists()
    with Image.open(FIG2_V67).convert("L") as im:
        w, h = im.size
        assert w >= 500 and h >= 400, "v67 crop 尺寸过小"
        ink = sum(1 for v in im.getdata() if v < 128) / (w * h)
        assert 0.02 <= ink <= 0.6, f"墨迹覆盖异常: {ink:.1%}"
        # hbs1 label zone (paper y 475-491, x 223-239) must carry ink
        px0 = int((223 - 62.8) * (w / (293.6 - 62.8)))
        px1 = int((239 - 62.8) * (w / (293.6 - 62.8)))
        py0 = int((475 - 376.8) * (h / (531.0 - 376.8)))
        py1 = int((491 - 376.8) * (h / (531.0 - 376.8)))
        dark = sum(1 for x in range(px0, px1, 2) for y in range(py0, py1, 2)
                   if im.getpixel((x, y)) < 128)
        tot = len(range(px0, px1, 2)) * len(range(py0, py1, 2))
        assert dark / tot > 0.02, "hbs1 标注区域缺失（尺寸标注被裁掉）"


# ── 4. generated_figure_no_internal_title_test ───────────────────

def test_generated_figure_no_internal_title_test(tmp_path: Path):
    """Fig-1/3/4 must NOT carry baked-in internal titles on the canvas."""
    figures = _plan_motor_figures()
    for number in (1, 3, 4):
        fig = next(f for f in figures if f.number == number)
        report = _render(fig, tmp_path)
        assert not report.get("title_bbox"), f"fig{number} 画布内部仍绘制标题: {report.get('title_bbox')}"
        kinds = {e["kind"] for e in report["elements"]}
        assert "title" not in kinds, f"fig{number} 存在 title 元素"
        # canvas top padding tightened: first node starts at the margin
        nodes = [e for e in report["elements"] if e["kind"] == "node"]
        assert nodes, f"fig{number} 无节点"
        assert min(n["bbox"][1] for n in nodes) < 100, \
            f"fig{number} 画布顶部 padding 未收紧（首节点 y={min(n['bbox'][1] for n in nodes)}）"


# ── 5. unexpected_blank_page_test ────────────────────────────────

def test_unexpected_blank_page_test(tmp_path: Path):
    """BlankPageValidator must flag content-free pages and pass clean ones."""
    import fitz

    # PDF with a genuinely blank page (page 2)
    bad_pdf = tmp_path / "blank.pdf"
    doc = fitz.open()
    p1 = doc.new_page()
    p1.insert_text((72, 72), "附图说明正文……")
    doc.new_page()  # blank
    p3 = doc.new_page()
    p3.insert_text((72, 72), "后续段落")
    doc.save(str(bad_pdf))
    doc.close()

    analysis = BlankPageValidator().analyze(bad_pdf)
    assert analysis["total_pages"] == 3
    assert analysis["blank_pages"] == [2], f"应检出第 2 页空白: {analysis['blank_pages']}"
    assert not analysis["pass"]
    findings = BlankPageValidator().validate(bad_pdf)
    assert [f.code for f in findings] == ["UNEXPECTED_BLANK_PAGE"]

    # PDF with content on every page -> pass
    good_pdf = tmp_path / "good.pdf"
    doc = fitz.open()
    for i in range(3):
        pg = doc.new_page()
        pg.insert_text((72, 72), f"第 {i + 1} 页正文")
    doc.save(str(good_pdf))
    doc.close()
    assert BlankPageValidator().analyze(good_pdf)["pass"]

    # page carrying ONLY a page-number footer is still a blank page
    footer_pdf = tmp_path / "footer_blank.pdf"
    doc = fitz.open()
    p1 = doc.new_page()
    p1.insert_text((72, 72), "正文内容……")
    p2 = doc.new_page()
    p2.insert_text((72, 780), "7")  # footer page number only
    doc.save(str(footer_pdf))
    doc.close()
    a = BlankPageValidator().analyze(footer_pdf)
    assert a["blank_pages"] == [2], "只有页码的页面必须判定为空白页"


# ── 6. no_forced_page_break_after_caption_test ───────────────────

def test_no_forced_page_break_after_caption_test():
    """Figures must flow dynamically: no forced page break nodes at all."""
    figures = [
        FigureSpec(id=f"FIG-{index:03d}", number=index, type="flowchart",
                   title=f"图{index}", nodes=[], edges=[], source_ids=[],
                   png_path=f"figure-{index}.png")
        for index in (1, 2)
    ]
    draft = DisclosureDraft(
        title="测试",
        sections={"8. 附图说明": []},
        equations=[],
        figures=figures,
        inventor_questions=[],
        evidence_ids=[],
    )
    ast = disclosure_to_ast("TEST", draft)
    breaks = [n for n in ast.nodes if n.type == "page_break"]
    assert breaks == [], "附图说明部分不得包含强制分页（V6.7 dynamic pagination）"


# ── 7. figure3_parameter_bridge_style_test ───────────────────────

def test_figure3_parameter_bridge_style_test(tmp_path: Path):
    """Fig-3 parameter bridge (T5->G2) must be DASHED, not solid."""
    from PIL import Image
    fig3 = next(f for f in _plan_motor_figures() if f.number == 3)
    report = _render(fig3, tmp_path)
    bridge = [e for e in report["elements"] if e["kind"] == "arrow" and "T5->G2" in e["node_id"]]
    assert len(bridge) == 3, "桥接箭头应有 3 个分段"
    labels = [e["content"] for e in report["elements"] if e["kind"] == "text" and "参数" in e["content"]]
    assert labels, "缺少训练→生成参数传递标签"

    # pixel proof: scan the bridge vertical segment for dash periodicity
    boxes = {e["node_id"]: e["bbox"] for e in report["elements"] if e["kind"] == "node"}
    t5, g2 = boxes["T5"], boxes["G2"]
    t5_x1, t5_y1 = t5[0] + t5[2], t5[1] + t5[3]
    g2_x0 = g2[0]
    s_cy, t_cy = (t5[1] + t5_y1) // 2, (g2[1] + g2[3]) // 2
    gap_mid_x = (t5_x1 + g2_x0) // 2
    im = Image.open(tmp_path / f"figure_{fig3.number:02d}.png").convert("L")
    lo, hi = min(s_cy, t_cy), max(s_cy, t_cy)
    dark = [y for y in range(lo, hi + 1) if im.getpixel((gap_mid_x, y)) < 128]
    runs, prev, start = 0, None, None
    for y in dark:
        if prev is None or y - prev > 3:
            runs += 1
            if runs == 2 and start is not None:
                pass
        prev = y
    n_runs = sum(1 for i, y in enumerate(dark) if i == 0 or y - dark[i - 1] > 3)
    assert n_runs >= 5, f"桥接竖段应为虚线（多个短 run），实际 {n_runs} 个 run"


# ── 8. figure4_latent_decoder_output_semantics_test ──────────────

def test_figure4_latent_decoder_output_semantics_test(tmp_path: Path):
    """Fig-4 must split latent Z / VAE decoder / output into 3 nodes."""
    fig4 = next(f for f in _plan_motor_figures() if f.number == 4)
    assert fig4.layout == "branch_merge"
    node_ids = {n.id for n in fig4.nodes}
    assert {"I1", "I2", "I3", "I4", "I5"} == node_ids, \
        f"图4 必须为 5 节点（Z1/Z2 + Z + 解码器 + 输出），实际 {node_ids}"
    labels = {n.id: n.label for n in fig4.nodes}
    assert "中间潜在变量" in labels["I3"], "I3 应为中间潜在变量 Z"
    assert "VAE解码器" in labels["I4"], "I4 应为 VAE 解码器"
    assert "平滑过渡拓扑序列" in labels["I5"], "I5 应为输出序列"
    edges = {(e.source, e.target) for e in fig4.edges}
    assert {"I1", "I2"} <= {s for s, _ in edges if _ == "I3"}, "Z1/Z2 必须合流到 I3"
    assert ("I3", "I4") in edges and ("I4", "I5") in edges, "I3→I4→I5 必须串联"
    report = _render(fig4, tmp_path)
    assert report["collisions"] == [], f"图4 存在碰撞: {report['collisions'][:5]}"


# ── 9. caption_consistency_test ──────────────────────────────────

def test_caption_consistency_test(tmp_path: Path):
    """Word captions: all "图N  <title>", centered, uniform spacing."""
    from PIL import Image
    image = tmp_path / "fig.png"
    Image.new("RGB", (400, 300), "white").save(image)
    ast = PatentDocumentAST(document_id="D", kind="disclosure", title="测试", nodes=[
        PatentNode(type="heading", value="附图说明", level=1),
        PatentNode(type="figure", target="FIG-001", path=str(image),
                   value="转子设计变量标注示意图（来源：原论文）", number=2),
    ])
    out = DocumentRenderer(tmp_path / "templates").render(ast, tmp_path / "out.docx")

    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    doc = Document(out)
    caps = [p for p in doc.paragraphs if p.text.startswith("图")]
    assert len(caps) == 1
    assert caps[0].text == "图2  转子设计变量标注示意图（来源：原论文）"
    assert caps[0].alignment == WD_ALIGN_PARAGRAPH.CENTER, "图题必须居中"
    assert caps[0].paragraph_format.keep_together is True
    # image paragraph: centered + keep with next caption
    blips = [p for p in doc.paragraphs
             if p.runs and any(r.element.findall(".//{http://schemas.openxmlformats.org/drawingml/2006/main}blip") for r in p.runs)]
    assert blips and blips[0].alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert blips[0].paragraph_format.keep_with_next is True
    # uniform spacing: caption style single source of truth (no duplicate runs)
    assert caps[0].runs and "图2" in caps[0].runs[0].text
