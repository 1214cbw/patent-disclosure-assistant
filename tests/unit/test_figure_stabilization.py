"""V6.6 disclosure figure stabilization - 10 named tests.

Covers: source-figure crop integrity / aspect ratio, figure-2 real-or-omit,
figure-3 two-column no-overlap, figure-4 branch-merge no-overlap, arrow-text
collision detection, Word caption-same-page, Word embed size bounds, semantic
validation and source provenance validation.
"""
from __future__ import annotations

import json
from pathlib import Path

import pytest

from patent_agent.agents.figure_planner import FigurePlanner
from patent_agent.core.models import FigureSpec, TechnicalUnderstandingResult
from patent_agent.core.patent_ast import PatentDocumentAST, PatentNode
from patent_agent.document import DocumentRenderer
from patent_agent.document.ast_factory import disclosure_to_ast
from patent_agent.document.figure_layout import (
    BBox,
    CollisionDetector,
    FigureLayoutValidator,
    FigureSemanticValidator,
    FigureSourceValidator,
    LayoutElement,
)
from patent_agent.document.figure_renderer import PatentFigureRenderer, _layout_report_path
from patent_agent.core.models import DisclosureDraft

EXTRACTED_FIG2 = Path("workspace/private_cases/REAL-PAPER-001/extracted_figures/fig2_design_variables.png")


def _plan_motor_figures() -> list[FigureSpec]:
    u = TechnicalUnderstandingResult.model_construct(
        steps=["motor rotor topology design", "latent diffusion training", "generation"],
        components=["motor", "rotor", "latent diffusion", "topology"],
    )
    return FigurePlanner().from_understanding(u, case_id="REAL-PAPER-001")


def _render(figure: FigureSpec, tmp_path: Path) -> dict:
    PatentFigureRenderer().render(figure, tmp_path)
    return json.loads(_layout_report_path(tmp_path, figure.number).read_text(encoding="utf-8"))


def test_vertical_feedback_edge_routes_outside_intermediate_nodes(tmp_path: Path):
    from patent_agent.core.models import FigureEdge, FigureNode

    figure = FigureSpec(
        id="FIG-FEEDBACK", number=9, type="flowchart", title="反馈流程",
        nodes=[FigureNode(id=f"N{i}", label=f"节点{i}") for i in range(1, 5)],
        edges=[
            FigureEdge(source="N1", target="N2"),
            FigureEdge(source="N2", target="N3"),
            FigureEdge(source="N3", target="N4"),
            FigureEdge(source="N4", target="N1"),
        ],
        source_ids=["F1"], layout="vertical",
    )
    report = _render(figure, tmp_path)
    assert report["collisions"] == []
    feedback_segments = [
        item for item in report["elements"]
        if item["kind"] == "arrow" and item["node_id"] == "N4->N1"
    ]
    assert len(feedback_segments) == 4


# ── 1. source_figure_crop_integrity_test ─────────────────────────

def test_source_figure_crop_integrity_test():
    """Extracted fig-2 must be a real crop: no black borders, sane ink."""
    from PIL import Image
    assert EXTRACTED_FIG2.exists(), "提取的真实图2必须存在"
    assert EXTRACTED_FIG2.stat().st_size > 10_000
    with Image.open(EXTRACTED_FIG2) as im:
        px = im.convert("L")
        w, h = px.size
        assert w >= 300 and h >= 300, "真实图裁剪尺寸过小"
        # black-band check: 2px border must not be uniformly black (93%黑垃圾裁切回归)
        for band in ("top", "bottom", "left", "right"):
            if band == "top":
                strip = [px.getpixel((x, 2)) for x in range(0, w, max(1, w // 40))]
            elif band == "bottom":
                strip = [px.getpixel((x, h - 3)) for x in range(0, w, max(1, w // 40))]
            elif band == "left":
                strip = [px.getpixel((2, y)) for y in range(0, h, max(1, h // 40))]
            else:
                strip = [px.getpixel((w - 3, y)) for y in range(0, h, max(1, h // 40))]
            black_frac = sum(1 for v in strip if v < 40) / len(strip)
            assert black_frac < 0.9, f"图2 {band} 边框存在大块黑边 (black={black_frac:.0%})"
        ink = sum(1 for v in px.getdata() if v < 128) / (w * h)
        assert 0.02 <= ink <= 0.6, f"墨迹覆盖异常: {ink:.1%}"


# ── 2. source_figure_aspect_ratio_test ───────────────────────────

def test_source_figure_aspect_ratio_test():
    """Embedded size must preserve the extracted figure's aspect ratio."""
    from PIL import Image
    # V6.7: source is the re-cropped v67 file (the planner's current png_path)
    fig = next(f for f in _plan_motor_figures() if f.number == 2)
    assert fig.provenance == "extracted" and fig.png_path
    with Image.open(fig.png_path) as im:
        src_aspect = im.height / im.width
    w_cm, h_cm = DocumentRenderer._figure_embed_size(fig.png_path)
    assert abs((h_cm / w_cm) - src_aspect) < 0.01, "嵌入尺寸破坏了宽高比"
    # renderer output must keep aspect too
    spec = PatentFigureRenderer().render(fig, Path("tmp/fig_smoke_v66"))
    with Image.open(spec.png_path) as out:
        assert abs((out.height / out.width) - src_aspect) < 0.01


# ── 3. figure2_real_structure_or_omit_test ───────────────────────

def test_figure2_real_structure_or_omit_test():
    """Fig-2 must be a real extracted figure or an explicit omission - never fake."""
    figures = _plan_motor_figures()
    fig2 = next(f for f in figures if f.number == 2)
    assert fig2.provenance in ("extracted", "omitted"), \
        f"图2来源必须是 extracted/omitted，当前: {fig2.provenance}"
    if fig2.provenance == "extracted":
        assert fig2.png_path and Path(fig2.png_path).exists()
        assert Path(fig2.png_path).stat().st_size > 10_000
    else:
        assert "待" in fig2.title, "omitted 图2标题必须标记待补充"


# ── 4. figure3_two_column_layout_no_overlap_test ─────────────────

def test_figure3_two_column_layout_no_overlap_test():
    """Fig-3 two-column training/generation must render with zero overlap."""
    fig3 = next(f for f in _plan_motor_figures() if f.number == 3)
    assert fig3.layout == "two_column"
    assert fig3.left_node_ids and fig3.right_node_ids
    report = _render(fig3, Path("tmp/fig_smoke_v66"))
    assert report["layout"] == "two_column"
    assert report["collisions"] == [], f"图3存在碰撞: {report['collisions'][:5]}"
    kinds = {e["kind"] for e in report["elements"]}
    assert {"node", "arrow", "text", "math"} <= kinds
    assert FigureLayoutValidator().validate(report) == []


# ── 5. figure4_branch_merge_layout_no_overlap_test ───────────────

def test_figure4_branch_merge_layout_no_overlap_test():
    """Fig-4 dual-input merge must render with zero overlap and both inputs."""
    fig4 = next(f for f in _plan_motor_figures() if f.number == 4)
    assert fig4.layout == "branch_merge"
    report = _render(fig4, Path("tmp/fig_smoke_v66"))
    assert report["layout"] == "branch_merge"
    assert report["collisions"] == [], f"图4存在碰撞: {report['collisions'][:5]}"
    node_ids = {e["node_id"] for e in report["elements"] if e["kind"] == "node"}
    assert {"I1", "I2", "I3", "I4", "I5"} <= node_ids, "图4缺少 I1-I5 节点（V6.7 语义拆分）"
    all_content = " ".join(e["content"] for e in report["elements"] if e["kind"] in ("text", "math"))
    assert "Z_1" in all_content or "Z1" in all_content, "缺少 Z1 输入"
    assert "Z_2" in all_content or "Z2" in all_content, "缺少 Z2 输入"
    assert "中间潜在变量" in all_content, "缺少中间潜在变量 Z 节点（V6.7 拆分）"
    assert "VAE解码器" in all_content, "缺少 VAE 解码器节点（V6.7 拆分）"
    assert "平滑过渡拓扑序列" in all_content, "缺少输出节点（V6.7 拆分）"
    assert "lambda" in all_content or "\\lambda" in all_content, "缺少插值公式"
    assert FigureLayoutValidator().validate(report) == []


# ── 6. arrow_text_collision_test ─────────────────────────────────

def test_arrow_text_collision_test():
    """Arrows may touch their own endpoints/label but never cross foreign text."""
    # synthetic: arrow crossing a foreign text must be reported
    detector = CollisionDetector(pad=2.0)
    elements = [
        LayoutElement("node", BBox(0, 0, 100, 40), node_id="N1"),
        LayoutElement("arrow", BBox(40, 20, 200, 6), node_id="A->B"),
        LayoutElement("text", BBox(80, 22, 40, 12), node_id="C", content="横穿文本"),
    ]
    types = {c.type for c in detector.detect(elements)}
    assert "arrow-node" in types and "arrow-text" in types
    # arrow's own label (same node_id) must not be flagged
    own = [
        LayoutElement("node", BBox(0, 0, 100, 40), node_id="A"),
        LayoutElement("arrow", BBox(50, 38, 6, 100), node_id="A->B"),
        LayoutElement("text", BBox(45, 60, 20, 10), node_id="A->B", content="参数传递"),
    ]
    assert detector.detect(own) == []
    # real fig-3: bridge label present and no arrow-text/arrow-node collisions
    fig3 = next(f for f in _plan_motor_figures() if f.number == 3)
    report = _render(fig3, Path("tmp/fig_smoke_v66"))
    labels = [e for e in report["elements"] if e["kind"] == "text" and "参数" in e["content"]]
    assert labels, "图3缺少训练→生成参数传递标签"
    assert not [c for c in report["collisions"] if c["type"].startswith("arrow")]


# ── 7. caption_same_page_test ────────────────────────────────────

def test_caption_same_page_test(tmp_path: Path):
    """Figure paragraph + caption must keep together / with next page."""
    from PIL import Image
    image = tmp_path / "fig.png"
    Image.new("RGB", (400, 300), "white").save(image)
    ast = PatentDocumentAST(document_id="D", kind="disclosure", title="测试", nodes=[
        PatentNode(type="heading", value="附图说明", level=1),
        PatentNode(type="figure", target="FIG-001", path=str(image), value="测试附图", number=1),
    ])
    out = DocumentRenderer(tmp_path / "templates").render(ast, tmp_path / "out.docx")
    from docx import Document
    doc = Document(out)
    image_pars = [p for p in doc.paragraphs if p.runs and any(r.element.findall(".//{http://schemas.openxmlformats.org/drawingml/2006/main}blip") for r in p.runs)]
    assert image_pars, "未找到图片段落"
    ip = image_pars[0]
    assert ip.paragraph_format.keep_with_next is True
    assert ip.paragraph_format.keep_together is True
    captions = [p for p in doc.paragraphs if p.text.startswith("图1")]
    assert captions, "未找到图题"
    assert captions[0].paragraph_format.keep_together is True


# ── 8. word_embedded_image_size_test ─────────────────────────────

def test_word_embedded_image_size_test(tmp_path: Path):
    """Embed size must respect 13.5cm width / 23cm height with aspect preserved."""
    from PIL import Image
    wide = tmp_path / "wide.png"; Image.new("RGB", (1000, 300), "white").save(wide)
    tall = tmp_path / "tall.png"; Image.new("RGB", (200, 1000), "white").save(tall)
    ww, wh = DocumentRenderer._figure_embed_size(str(wide))
    tw, th = DocumentRenderer._figure_embed_size(str(tall))
    assert ww.cm == pytest.approx(13.5) and wh.cm <= 23.0
    assert th.cm == pytest.approx(23.0) and tw.cm <= 13.5
    assert (wh.cm / ww.cm) == pytest.approx(300 / 1000, abs=0.01)
    assert (th.cm / tw.cm) == pytest.approx(1000 / 200, abs=0.01)


# ── 9. figure_semantic_validation_test ───────────────────────────

def test_figure_semantic_validation_test():
    figures = _plan_motor_figures()
    v = FigureSemanticValidator()
    fig3 = next(f for f in figures if f.number == 3)
    issues3 = v.validate(fig3)
    codes3 = {i.code for i in issues3}
    assert not codes3 & {"NO_TRAINING_PATH", "NO_GENERATION_PATH", "NO_PARAMETER_BRIDGE"}, issues3
    fig4 = next(f for f in figures if f.number == 4)
    codes4 = {i.code for i in v.validate(fig4)}
    assert not codes4 & {"NO_DUAL_INPUT", "NO_MERGE_NODE", "NO_MERGE_OUTPUT"}, issues3
    # fake figure-2 (provenance=generated) must be rejected
    fake2 = next(f for f in figures if f.number == 2).model_copy(update={"provenance": "generated"})
    assert {i.code for i in v.validate(fake2)} & {"FAKE_STRUCTURE_FIGURE"}
    # broken figure-3 (no generation path) must be caught
    broken3 = fig3.model_copy(update={"nodes": [n for n in fig3.nodes if n.id in fig3.left_node_ids]})
    assert {i.code for i in v.validate(broken3)} & {"NO_GENERATION_PATH"}


# ── 10. figure_source_provenance_test ────────────────────────────

def test_figure_source_provenance_test():
    v = FigureSourceValidator()
    figures = _plan_motor_figures()
    assert v.validate(figures) == []
    bad = [f.model_copy(update={"provenance": "pasted"}) for f in figures]
    assert {i.code for i in v.validate(bad)} == {"UNKNOWN_PROVENANCE"}
    no_img = figures[0].model_copy(update={"provenance": "extracted", "png_path": ""})
    assert {i.code for i in v.validate([no_img])} == {"EXTRACTED_WITHOUT_IMAGE"}
